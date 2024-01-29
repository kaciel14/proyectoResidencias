from docx import Document
from datetime import datetime

class ForDocx:

    def __init__(self, data, params, ruta):

        now = datetime.now()
        self.rutaOut = './../archivos/' + str(now.time()).replace(':', '-') + '.docx'

        self.rutaIn = ruta.strip()

        self.document = Document(self.rutaIn)

        self.document.save(self.rutaIn)

        self.document = Document(self.rutaIn)

        #paragraph = document.add_paragraph('KACIEL BENITEZ')

        #PARAMETROS QUE SE OBTENDRAN DE LA BASE DE DATOS DEPENDIENDO DE LA PLANTILLA SELECCIONADA
        #params = ['NOMBRE', 'TITULO', 'NUMCONTROL', 'COLONIA', 'CIUDAD', 'TELEFONO']

        #self.params = ['NOMBRE']

        self.params = []

        params = params.split(', ')
        
        self.params = params

        #self.params.append(params)

        #VALORES DE ENTRADA RECIBIDOS DEL SERVIDOR PASADOS POR EL BOT (CORRESPONDEN A LOS PARAMETROS DE LA BASE DE DATOS) (EL ORDEN IMPORTA)
        #inputs = ['Kaciel Alejandro Benitez Ferral', 'PRIMER REPORTE', '19071482', 'Col. Los Mangos', 'Tampico, Tamaulipas', '833-357-48-20']

        #inputs = []

        data = data.split(',')

        self.inputs = data

        text = self.document.paragraphs
                            
        tables = self.document.tables


        if(text != [] and tables != []):
            self.readParagraphs(text)
            self.readTables(tables)
        elif(text != []):
            self.readParagraphs(text)
        elif(tables != []):
            self.readTables(tables)

        self.document.save(self.rutaOut)


    def readParagraphs(self, text):
        for para in text:
            for param, input in zip(self.params, self.inputs):
                #print(param)
                if('['+ param +']' in para.text):
                    
                    for run in para.runs:
                        if '[' + param + ']' in run.text:
                            #print('A: ' + run.text)
                            run.text = run.text.replace('['+ param +']', input)
                        #else:
                            #print('B: '+run.text)        
                    

                            

                        #print("remplazado: " + '[['+param+']]'+ " por: " + run.text)
                    #para.text = para.text.replace('['+ param +']', input)

    def readTables(self, tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for param, input in zip(self.params, self.inputs):
                            if('['+ param +']' in para.text):
                                for run in para.runs:
                                    if '[' + param + ']' in run.text:
                                        #print('A: ' + run.text)
                                        run.text = run.text.replace('['+ param +']', input)

    def getRutaOut(self):
        return self.rutaOut[2:]
    